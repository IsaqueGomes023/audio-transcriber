"""
exporter.py
-----------
Módulo de exportação dos resultados (transcrição + resumo) para TXT e DOCX.

Funcionalidades:
  - Exportar para TXT simples (encoding UTF-8)
  - Exportar para DOCX formatado com estilos profissionais
  - Gerar bytes prontos para download via Streamlit

Dependências: python-docx
"""

import io
import re
import logging
from datetime import datetime
from typing import Optional

logger = logging.getLogger(__name__)


# ─── Exportação TXT ──────────────────────────────────────────────────────────

def to_txt(
    transcription: str,
    summary: str,
    audio_filename: str = "audio",
    language: str = "pt",
    duration: str = "desconhecida",
    include_segments: Optional[str] = None,
) -> bytes:
    """
    Gera um arquivo TXT completo com transcrição e resumo.

    Args:
        transcription:    Texto completo da transcrição.
        summary:          Texto do resumo (Markdown).
        audio_filename:   Nome do arquivo de áudio original.
        language:         Idioma detectado.
        duration:         Duração do áudio.
        include_segments: Transcrição segmentada com timestamps (opcional).

    Returns:
        Conteúdo do arquivo TXT como bytes (UTF-8).
    """
    now = datetime.now().strftime("%d/%m/%Y %H:%M")
    separator = "=" * 70

    lines = [
        separator,
        "  TRANSCRIÇÃO E RESUMO DE ÁUDIO",
        f"  Arquivo: {audio_filename}",
        f"  Data: {now}",
        f"  Idioma: {language.upper()} | Duração: {duration}",
        separator,
        "",
        separator,
        "  RESUMO INTELIGENTE",
        separator,
        "",
        _strip_markdown(summary),
        "",
        separator,
        "  TRANSCRIÇÃO COMPLETA",
        separator,
        "",
        transcription,
        "",
    ]

    if include_segments:
        lines += [
            separator,
            "  TRANSCRIÇÃO COM TIMESTAMPS",
            separator,
            "",
            include_segments,
            "",
        ]

    lines += [
        separator,
        f"  Gerado por AudioTranscriber — {now}",
        separator,
    ]

    content = "\n".join(lines)
    return content.encode("utf-8")


def _strip_markdown(text: str) -> str:
    """Remove marcações Markdown para saída em texto puro."""
    # Remove headers (##, ###, etc.)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Remove bold/italic (**texto**, *texto*)
    text = re.sub(r"\*{1,3}(.+?)\*{1,3}", r"\1", text)
    # Remove links [texto](url)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Remove checkbox [ ] e [x]
    text = re.sub(r"\[[ xX]\]\s*", "• ", text)
    # Remove linhas horizontais (---)
    text = re.sub(r"^---+\s*$", "-" * 50, text, flags=re.MULTILINE)
    return text.strip()


# ─── Exportação DOCX ─────────────────────────────────────────────────────────

def to_docx(
    transcription: str,
    summary: str,
    audio_filename: str = "audio",
    language: str = "pt",
    duration: str = "desconhecida",
    include_segments: Optional[str] = None,
) -> bytes:
    """
    Gera um arquivo DOCX formatado profissionalmente.

    Args:
        (mesmos parâmetros de to_txt)

    Returns:
        Conteúdo do arquivo DOCX como bytes.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx não instalado. Execute: pip install python-docx"
        )

    doc = Document()
    now = datetime.now().strftime("%d/%m/%Y %H:%M")

    # ── Configuração da página ──
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Estilos base ──
    styles = doc.styles

    # ── Cabeçalho do documento ──
    title = doc.add_heading("Transcrição e Resumo de Áudio", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8A)  # Azul escuro

    # Metadados
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(meta, f"📁 Arquivo: ", bold=True)
    _add_run(meta, audio_filename)
    meta.add_run("   |   ")
    _add_run(meta, "📅 Data: ", bold=True)
    _add_run(meta, now)
    meta.add_run("   |   ")
    _add_run(meta, "🌐 Idioma: ", bold=True)
    _add_run(meta, language.upper())
    meta.add_run("   |   ")
    _add_run(meta, "⏱ Duração: ", bold=True)
    _add_run(meta, duration)

    doc.add_paragraph()  # Espaço

    # ── Seção: Resumo ──
    h1 = doc.add_heading("Resumo Inteligente", 1)
    h1.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8A)

    _render_markdown_to_docx(doc, summary)

    doc.add_page_break()

    # ── Seção: Transcrição completa ──
    h2 = doc.add_heading("Transcrição Completa", 1)
    h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8A)

    trans_para = doc.add_paragraph(transcription)
    trans_para.runs[0].font.size = Pt(11)
    trans_para.paragraph_format.space_after = Pt(6)

    # ── Seção: Timestamps (opcional) ──
    if include_segments:
        doc.add_page_break()
        h3 = doc.add_heading("Transcrição com Timestamps", 1)
        h3.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8A)

        for line in include_segments.split("\n"):
            if line.strip():
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.name = "Courier New"

    # ── Rodapé ──
    _add_footer(doc, now)

    # Salva em buffer de memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_run(paragraph, text: str, bold: bool = False, italic: bool = False,
             size_pt: int = 11, color: Optional[RGBColor] = None):
    """Adiciona um run formatado a um parágrafo."""
    from docx.shared import Pt, RGBColor as RC
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def _render_markdown_to_docx(doc, markdown_text: str):
    """
    Converte Markdown básico em elementos DOCX estruturados.
    Suporta: ## headers, - bullets, **bold**, [ ] checkboxes, --- separadores.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Linha em branco
        if not stripped:
            i += 1
            continue

        # Separador horizontal
        if stripped.startswith("---"):
            doc.add_paragraph("─" * 60).paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Cabeçalhos ## e ###
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:].strip(), level=3)
            h.runs[0].font.size = Pt(12)
            i += 1
            continue

        if stripped.startswith("## "):
            h = doc.add_heading(stripped[3:].strip(), level=2)
            h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            i += 1
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:].strip(), level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x45, 0x8A)
            i += 1
            continue

        # Checkboxes [ ] e [x]
        if stripped.startswith("- [ ]") or stripped.startswith("- [x]") or stripped.startswith("- [X]"):
            checked = stripped[3] in ("x", "X")
            text = stripped[6:].strip()
            p = doc.add_paragraph(style="List Bullet")
            prefix = "☑ " if checked else "☐ "
            run = p.add_run(prefix + text)
            run.font.size = Pt(11)
            i += 1
            continue

        # Bullets (-  ou *)
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)
            i += 1
            continue

        # Parágrafo normal (pode conter **bold**)
        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)
        i += 1


def _add_formatted_run(paragraph, text: str):
    """Processa **bold** e *italic* inline dentro de um parágrafo."""
    from docx.shared import Pt
    import re

    # Padrão: **texto** ou *texto*
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*)")
    last_end = 0

    for match in pattern.finditer(text):
        # Texto antes do match
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.size = Pt(11)

        matched = match.group(0)
        if matched.startswith("**"):
            run = paragraph.add_run(match.group(2))
            run.bold = True
        else:
            run = paragraph.add_run(match.group(3))
            run.italic = True
        run.font.size = Pt(11)
        last_end = match.end()

    # Texto restante
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(11)


def _add_footer(doc, date_str: str):
    """Adiciona rodapé ao documento."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = doc.sections[0]
    footer  = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Gerado por AudioTranscriber | {date_str}")
    run.font.size = Pt(9)
    run.font.italic = True


# ─── Helpers para Streamlit ──────────────────────────────────────────────────

def get_download_filename(audio_filename: str, ext: str) -> str:
    """Gera nome de arquivo para download com timestamp."""
    base = re.sub(r"[^\w\-_]", "_", audio_filename.rsplit(".", 1)[0])
    now  = datetime.now().strftime("%Y%m%d_%H%M")
    return f"transcricao_{base}_{now}.{ext}"
